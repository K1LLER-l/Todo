from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def crear_manual():
    doc = Document()

    # --- ESTILOS ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- PORTADA ---
    titulo = doc.add_heading('Manual Técnico y Guía de Código', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitulo = doc.add_paragraph('Proyecto: Todo Protein')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.style = 'Subtitle'
    
    doc.add_paragraph(f'\nFecha de generación: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph('\nEste documento detalla la estructura, archivos y funcionamiento lógico del sistema Todo Protein.')
    doc.add_page_break()

    # --- SECCIÓN 1: ESTRUCTURA GENERAL ---
    doc.add_heading('1. Estructura del Proyecto Django', level=1)
    doc.add_paragraph('El proyecto sigue la arquitectura MVT (Modelo-Vista-Template) clásica de Django.')
    
    estructura = """
    TodoProtein/
    ├── catalog/               <-- Aplicación principal (Lógica de negocio)
    │   ├── management/        <-- Scripts de automatización (Scraping, Alertas)
    │   ├── migrations/        <-- Historial de cambios en la Base de Datos
    │   ├── templatetags/      <-- Filtros personalizados para HTML
    │   ├── admin.py           <-- Configuración del panel de control
    │   ├── models.py          <-- Definición de las tablas de BD
    │   ├── views.py           <-- Lógica que conecta datos con pantallas
    │   └── urls.py            <-- (Opcional) Rutas específicas de la app
    ├── loginProyecto/         <-- Configuración global del sitio
    │   ├── settings.py        <-- Configuración (BD, Apps, Correos, Seguridad)
    │   └── urls.py            <-- Tabla de contenidos (Rutas) principal
    ├── static/                <-- Archivos estáticos (CSS, JS, Imágenes)
    ├── templates/             <-- Archivos HTML (El diseño visual)
    └── manage.py              <-- El cerebro para ejecutar comandos
    """
    p = doc.add_paragraph(estructura)
    p.style = 'No Spacing'
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    # --- SECCIÓN 2: DICCIONARIO DE ARCHIVOS ---
    doc.add_heading('2. Diccionario de Archivos Clave', level=1)
    
    archivos = [
        ("catalog/models.py", "El Corazón de los Datos", 
         "Aquí se definen las 'Tablas' de la base de datos como clases de Python. \n"
         "- Producto: Nombre, marca, info nutricional.\n"
         "- Oferta: Precio y enlace a tienda.\n"
         "- PrecioHistorico: Registro de precios pasados.\n"
         "- Favorito: Relación entre Usuario y Producto con precio deseado."),
        
        ("catalog/views.py", "El Cerebro Lógico", 
         "Aquí están las funciones que deciden qué mostrar al usuario.\n"
         "- home(): Busca productos, filtra y muestra el ranking.\n"
         "- product_detail(): Muestra el gráfico, la tabla nutricional y procesa las alertas.\n"
         "- set_precio_deseado(): Guarda la meta de precio del usuario."),
        
        ("catalog/admin.py", "El Panel de Control", 
         "Configura cómo se ve el administrador de Django (/admin). Aquí activamos la edición de Tiendas y configuración de Scraping."),
        
        ("loginProyecto/urls.py", "El Mapa de Navegación", 
         "Define las rutas URL. Dice: 'Si el usuario entra a /favoritos, ejecuta la vista lista_favoritos'."),
        
        ("loginProyecto/settings.py", "La Configuración Global", 
         "Archivo crítico. Contiene las credenciales de base de datos, claves secretas, configuración de correo Gmail y validadores de contraseña."),
        
        ("catalog/templatetags/mis_filtros.py", "Herramientas Visuales", 
         "Contiene filtros personalizados como 'redondear' o 'precio' para dar formato bonito a los números en el HTML.")
    ]

    for archivo, apodo, desc in archivos:
        h = doc.add_heading(archivo, level=2)
        h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        p = doc.add_paragraph()
        p.add_run(f"Apodo: {apodo}\n").bold = True
        p.add_run(desc)

    # --- SECCIÓN 3: SCRIPTS Y COMANDOS ---
    doc.add_page_break()
    doc.add_heading('3. Comandos de Automatización (Management Commands)', level=1)
    doc.add_paragraph('Estos scripts se ejecutan desde la terminal y realizan tareas en segundo plano.')

    comandos = [
        ("python manage.py scrape_category_pages", "El Robot Buscador",
         "Navega por las webs configuradas, extrae precios y productos nuevos, y los guarda en la base de datos."),
        
        ("python manage.py enviar_alerta", "El Notificador Instantáneo",
         "Revisa si algún producto favorito bajó de precio HOY y envía un correo de alerta si se cumple la meta del usuario."),
        
        ("python manage.py enviar_resumen_semanal", "El Reporte Semanal",
         "Compara precios de hoy vs. hace 7 días y envía un resumen consolidado por correo."),
        
        ("python manage.py iniciar_scheduler", "El Piloto Automático",
         "Mantiene el sistema vivo. Ejecuta el scraper y las alertas automáticamente según el horario programado.")
    ]

    for cmd, nombre, desc in comandos:
        p = doc.add_paragraph()
        run = p.add_run(f"► {nombre}")
        run.bold = True
        run.font.color.rgb = RGBColor(230, 126, 34) # Naranja
        doc.add_paragraph(f"Comando: {cmd}", style='Quote')
        doc.add_paragraph(desc)

    # --- SECCIÓN 4: TEMPLATES (FRONTEND) ---
    doc.add_heading('4. Plantillas HTML (Lo que ve el usuario)', level=1)
    
    templates = [
        ("base.html", "La Plantilla Madre. Contiene el menú (navbar), el pie de página (footer) y los estilos CSS comunes. Todas las demás páginas 'heredan' de esta."),
        ("home.html", "La Página de Inicio. Muestra el buscador, el banner de promociones y la grilla de productos."),
        ("product_detail.html", "La Ficha del Producto. Muestra el gráfico de precios, la tabla nutricional y el formulario de alertas."),
        ("comparison.html", "La Tabla Comparativa. Muestra productos lado a lado con sus calorías y precios."),
        ("favorites.html", "Tu Panel Personal. Tabla resumen con tus productos guardados y estado de metas.")
    ]

    for temp, desc in templates:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(temp).bold = True
        p.add_run(f": {desc}")

    # --- Guardar ---
    nombre_archivo = 'Manual_Tecnico_TodoProtein.docx'
    doc.save(nombre_archivo)
    print(f"¡Manual generado exitosamente: {nombre_archivo}!")

if __name__ == "__main__":
    crear_manual()